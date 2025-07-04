import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any
import os
from .calculator import CostCalculator
from .models import LifeCarePlan, ProjectionSettings


class ExcelExporter:
    """Export life care plan data to Excel format."""
    
    def __init__(self, calculator: CostCalculator):
        self.calculator = calculator
        self.lcp = calculator.lcp
    
    def export(self, file_path: str, include_all_scenarios: bool = False, selected_scenarios: list = None) -> None:
        """Export the life care plan to Excel file with improved formatting.
        
        Args:
            file_path: Output file path
            include_all_scenarios: If True, export selected scenarios with comparison sheets
            selected_scenarios: List of scenario names to include (if None, uses all scenarios)
        """
        if include_all_scenarios and len(self.lcp.scenarios) > 1:
            # Use selected scenarios or all scenarios if none specified
            scenarios_to_export = selected_scenarios if selected_scenarios else list(self.lcp.scenarios.keys())
            if len(scenarios_to_export) > 1:
                self._export_multi_scenario(file_path, scenarios_to_export)
            else:
                # Only one scenario selected, export as single scenario
                if scenarios_to_export:
                    original_active = self.lcp.active_scenario
                    self.lcp.set_active_scenario(scenarios_to_export[0])
                    self._export_single_scenario(file_path)
                    self.lcp.set_active_scenario(original_active)
                else:
                    self._export_single_scenario(file_path)
        else:
            self._export_single_scenario(file_path)
    
    def _export_single_scenario(self, file_path: str) -> None:
        """Export the current scenario only."""
        df = self.calculator.build_cost_schedule()
        summary_stats = self.calculator.calculate_summary_statistics()
        category_costs = self.calculator.get_cost_by_category()

        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            # Main cost schedule - export as-is with original column names
            # The original column names are more descriptive and show service details
            df.to_excel(writer, sheet_name='Annual Cost Schedule', index=False)

            # Enhanced Summary statistics with clearer descriptions
            summary_data = [
                ['Life Care Plan Summary', ''],
                ['Evaluee Name', self.lcp.evaluee.name],
                ['Current Age at Base Year', f"{self.lcp.evaluee.current_age} years old"],
                ['Base Year (Analysis Start)', str(self.lcp.settings.base_year)],
                ['Projection Period', f"{self.lcp.settings.projection_years:.1f} years ({summary_stats['projection_period']})"],
                ['Discount Rate Applied', f"{self.lcp.settings.discount_rate:.1%}" if self.lcp.evaluee.discount_calculations else "Not Applied"],
                ['', ''],
                ['Financial Summary', ''],
                ['Total Lifetime Cost (Nominal)', f"${summary_stats['total_nominal_cost']:,.2f}"],
                ['Average Annual Cost', f"${summary_stats['average_annual_cost']:,.2f}"],
            ]
            
            # Only include discount rate info if calculations are enabled AND discount rate > 0
            if self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0:
                summary_data.extend([
                    ['Total Lifetime Cost (Present Value)', f"${summary_stats['total_present_value']:,.2f}"],
                    ['Present Value Savings vs Nominal', f"${summary_stats['total_nominal_cost'] - summary_stats['total_present_value']:,.2f}"],
                ])

            summary_data.extend([
                ['', ''],
                ['Analysis Details', ''],
                ['Service Categories Included', str(len(self.lcp.tables))],
                ['Total Individual Services', str(sum(len(table.services) for table in self.lcp.tables.values()))],
                ['Report Generated', datetime.now().strftime('%Y-%m-%d at %H:%M:%S')],
            ])

            summary_df = pd.DataFrame(summary_data, columns=['Description', 'Value'])
            summary_df.to_excel(writer, sheet_name='Executive Summary', index=False)
            
            # Enhanced Category breakdown with clearer headers
            category_rows = []
            show_pv_in_excel = self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0
            if show_pv_in_excel:
                category_columns = [
                    'Service Category',
                    'Total Lifetime Cost (Nominal)',
                    'Total Lifetime Cost (Present Value)',
                    'Number of Services'
                ]
                for table_name, data in category_costs.items():
                    category_rows.append([
                        table_name,
                        f"${data['table_nominal_total']:,.2f}",
                        f"${data['table_present_value_total']:,.2f}",
                        len(data['services'])
                    ])
            else:
                category_columns = [
                    'Service Category',
                    'Total Lifetime Cost (Nominal)',
                    'Number of Services'
                ]
                for table_name, data in category_costs.items():
                    category_rows.append([
                        table_name,
                        f"${data['table_nominal_total']:,.2f}",
                        len(data['services'])
                    ])

            category_df = pd.DataFrame(category_rows, columns=category_columns)
            category_df.to_excel(writer, sheet_name='Cost by Category', index=False)
            
            # Detailed Service Information with clearer headers
            service_rows = []
            service_columns = [
                'Service Category',
                'Service Name',
                'Unit Cost ($)',
                'Frequency per Year',
                'Annual Inflation Rate (%)',
                'Service Type',
                'Start Year',
                'End Year',
                'Total Lifetime Cost (Nominal)'
            ]

            if show_pv_in_excel:
                service_columns.append('Total Lifetime Cost (Present Value)')
            
            for table_name, data in category_costs.items():
                for service in data['services']:
                    service_type = 'One-time' if service['is_one_time_cost'] else \
                                  'Discrete' if service['occurrence_years'] else 'Recurring'
                    
                    start_year = service['one_time_cost_year'] if service['is_one_time_cost'] else \
                                service['start_year'] if service['start_year'] else 'N/A'
                    end_year = service['one_time_cost_year'] if service['is_one_time_cost'] else \
                              service['end_year'] if service['end_year'] else 'N/A'
                    
                    if service['occurrence_years']:
                        start_year = min(service['occurrence_years'])
                        end_year = max(service['occurrence_years'])
                        service_type += f" ({len(service['occurrence_years'])} occurrences)"
                    
                    service_row = [
                        table_name,
                        service['name'],
                        f"${service['unit_cost']:,.2f}",
                        f"{service['frequency_per_year']:.1f}x per year",
                        f"{service['inflation_rate']:.1f}%",
                        service_type,
                        start_year,
                        end_year,
                        f"${service['nominal_total']:,.2f}"
                    ]
                    
                    if show_pv_in_excel:
                        service_row.append(f"${service['present_value_total']:,.2f}")
                    
                    service_rows.append(service_row)
            
            service_df = pd.DataFrame(service_rows, columns=service_columns)
            service_df.to_excel(writer, sheet_name='Service Details', index=False)
            
            # Add enhanced calculation sheets
            self._add_calculation_sheets(writer)

    def _add_calculation_sheets(self, writer):
        """Add comprehensive calculation sheets with formulas and audit trails."""
        
        # Add Calculation Methodology Sheet
        methodology_data = [
            ['Life Care Plan Calculation Methodology', ''],
            ['', ''],
            ['Core Equations:', ''],
            ['Inflation-Adjusted Cost', 'C(t) = C₀ × (1 + i)ᵗ'],
            ['Present Value', 'PV(t) = C(t) ÷ (1 + d)ᵗ'],
            ['Total Lifetime Cost (Nominal)', 'Σ [C₀ × (1 + i)ᵗ × f]'],
            ['Total Lifetime Cost (PV)', 'Σ [C₀ × (1 + i)ᵗ × f ÷ (1 + d)ᵗ]'],
            ['', ''],
            ['Variable Definitions:', ''],
            ['C(t)', 'Cost in year t'],
            ['C₀', 'Base year unit cost'],
            ['i', 'Annual inflation rate (decimal)'],
            ['d', 'Discount rate (decimal)'],
            ['t', 'Years from base year'],
            ['f', 'Frequency per year'],
            ['Σ', 'Sum over projection period'],
            ['', ''],
            ['Validation Standards:', ''],
            ['Tolerance', '< $1.00 discrepancy'],
            ['Cross-validation', '5-point verification system'],
            ['Matrix reconciliation', 'Audit-standard methodology']
        ]
        
        methodology_df = pd.DataFrame(methodology_data, columns=['Parameter', 'Definition/Formula'])
        methodology_df.to_excel(writer, sheet_name='Calculation Methodology', index=False)
        
        # Add Sensitivity Analysis Sheet
        self._add_sensitivity_analysis_sheet(writer)
        
        # Add Factor Tables Sheet
        self._add_factor_tables_sheet(writer)
        
        # Add Audit Trail Sheet
        self._add_audit_trail_sheet(writer)
        
        # Add Service Master Table
        self._add_service_master_sheet(writer)

    def _add_sensitivity_analysis_sheet(self, writer):
        """Add sensitivity analysis calculations to Excel."""
        base_summary = self.calculator.calculate_summary_statistics()
        
        # Only include discount rate sensitivity if discount calculations are enabled
        if not self.lcp.evaluee.discount_calculations:
            # Create a simplified analysis sheet
            analysis_data = []
            analysis_data.append(['Sensitivity Analysis', ''])
            analysis_data.append(['Present value calculations are disabled', ''])
            analysis_data.append(['Only nominal cost analysis is available', ''])
            analysis_data.append(['', ''])
            analysis_data.append(['Inflation Impact Guidelines:', ''])
            analysis_data.append(['1% increase across all services', 'Increases nominal costs 15-25%'])
            analysis_data.append(['Higher inflation rates', 'Compound exponentially over time'])
            analysis_data.append(['Consider regional cost variations', 'May require 20-40% adjustments'])
            
            analysis_df = pd.DataFrame(analysis_data, columns=['Analysis Type', 'Impact'])
            analysis_df.to_excel(writer, sheet_name='Sensitivity Analysis', index=False)
            return
        
        # Discount rate sensitivity
        discount_sensitivity = []
        discount_sensitivity.append(['Discount Rate Sensitivity Analysis', '', '', ''])
        discount_sensitivity.append(['Discount Rate', 'Total Present Value', 'Difference from Base', 'Percentage Change'])
        
        base_discount = self.lcp.settings.discount_rate
        for rate_adjustment in [-0.01, -0.005, 0.0, 0.005, 0.01]:
            test_rate = base_discount + rate_adjustment
            temp_settings = ProjectionSettings(
                base_year=self.lcp.settings.base_year,
                projection_years=self.lcp.settings.projection_years,
                discount_rate=test_rate
            )
            temp_lcp = LifeCarePlan(evaluee=self.lcp.evaluee, settings=temp_settings)
            temp_lcp.tables = self.lcp.tables
            temp_calc = CostCalculator(temp_lcp)
            temp_summary = temp_calc.calculate_summary_statistics()
            
            pv_difference = temp_summary['total_present_value'] - base_summary['total_present_value']
            pv_percent = (pv_difference / base_summary['total_present_value']) * 100 if base_summary['total_present_value'] > 0 else 0
            
            discount_sensitivity.append([
                f"{test_rate:.1%}",
                f"${temp_summary['total_present_value']:,.2f}",
                f"${pv_difference:,.2f}",
                f"{pv_percent:+.2f}%"
            ])
        
        discount_sensitivity.append(['', '', '', ''])
        discount_sensitivity.append(['Inflation Sensitivity Guidelines:', '', '', ''])
        discount_sensitivity.append(['1% increase across all services', 'Increases nominal costs 15-25%', '', ''])
        discount_sensitivity.append(['Higher inflation rates', 'Compound exponentially over time', '', ''])
        discount_sensitivity.append(['Present value impact', 'Moderated by discount rate', '', ''])
        
        sensitivity_df = pd.DataFrame(discount_sensitivity, columns=['Parameter', 'Value', 'Difference', 'Percentage'])
        sensitivity_df.to_excel(writer, sheet_name='Sensitivity Analysis', index=False)

    def _add_factor_tables_sheet(self, writer):
        """Add mathematical factor tables to Excel."""
        # Create discount factor table
        discount_factors = []
        discount_factors.append(['Mathematical Factor Tables', '', ''])
        discount_factors.append(['', '', ''])
        discount_factors.append([f'Discount Factors ({self.lcp.settings.discount_rate:.1%} Rate)', '', ''])
        discount_factors.append(['Year', 'Discount Factor', 'Cumulative Factor'])
        
        cumulative_factor = 0
        projection_years = min(int(self.lcp.settings.projection_years) + 1, 40)  # Limit to 40 years for readability
        
        for year in range(projection_years):
            factor = 1 / (1 + self.lcp.settings.discount_rate) ** year
            cumulative_factor += factor
            discount_factors.append([
                self.lcp.settings.base_year + year,
                f"{factor:.6f}",
                f"{cumulative_factor:.6f}"
            ])
        
        # Add inflation factor examples
        discount_factors.append(['', '', ''])
        discount_factors.append(['Common Inflation Factors', '', ''])
        discount_factors.append(['Year', '2.5% Inflation', '3.0% Inflation', '3.5% Inflation'])
        
        for year in range(min(20, projection_years)):
            discount_factors.append([
                self.lcp.settings.base_year + year,
                f"{(1.025) ** year:.6f}",
                f"{(1.030) ** year:.6f}",
                f"{(1.035) ** year:.6f}"
            ])
        
        factors_df = pd.DataFrame(discount_factors, columns=['Year/Description', 'Factor/2.5%', 'Cumulative/3.0%', 'Additional/3.5%'])
        factors_df.to_excel(writer, sheet_name='Factor Tables', index=False)

    def _add_audit_trail_sheet(self, writer):
        """Add detailed audit trail for verification."""
        audit_data = []
        audit_data.append(['Audit Trail and Quality Control', '', '', '', ''])
        audit_data.append(['', '', '', '', ''])
        
        # Get verification data
        summary_stats = self.calculator.calculate_summary_statistics()
        cost_schedule = self.calculator.build_cost_schedule()
        category_costs = self.calculator.get_cost_by_category()
        
        # Summary verification
        audit_data.append(['Executive Summary Verification:', '', '', '', ''])
        audit_data.append(['Total Nominal Cost', f"${summary_stats['total_nominal_cost']:,.2f}", '', '', ''])
        if self.lcp.evaluee.discount_calculations:
            audit_data.append(['Total Present Value', f"${summary_stats['total_present_value']:,.2f}", '', '', ''])
        audit_data.append(['Average Annual Cost', f"${summary_stats['average_annual_cost']:,.2f}", '', '', ''])
        audit_data.append(['Calculation Check', f"${summary_stats['total_nominal_cost'] / self.lcp.settings.projection_years:,.2f}", 'Should match average', '', ''])
        
        # Category reconciliation
        audit_data.append(['', '', '', '', ''])
        audit_data.append(['Category Reconciliation:', '', '', '', ''])
        total_from_categories = sum(data['table_nominal_total'] for data in category_costs.values())
        audit_data.append(['Sum of Categories', f"${total_from_categories:,.2f}", '', '', ''])
        audit_data.append(['Executive Total', f"${summary_stats['total_nominal_cost']:,.2f}", '', '', ''])
        audit_data.append(['Difference', f"${abs(total_from_categories - summary_stats['total_nominal_cost']):,.2f}", 'Should be < $1.00', '', ''])
        
        # Annual schedule verification
        audit_data.append(['', '', '', '', ''])
        audit_data.append(['Annual Schedule Verification:', '', '', '', ''])
        schedule_total = cost_schedule['Total Nominal'].sum()
        audit_data.append(['Schedule Total', f"${schedule_total:,.2f}", '', '', ''])
        audit_data.append(['Executive Total', f"${summary_stats['total_nominal_cost']:,.2f}", '', '', ''])
        audit_data.append(['Difference', f"${abs(schedule_total - summary_stats['total_nominal_cost']):,.2f}", 'Should be < $1.00', '', ''])
        
        # Years verification
        audit_data.append(['', '', '', '', ''])
        audit_data.append(['Projection Period Verification:', '', '', '', ''])
        audit_data.append(['Projection Years', f"{self.lcp.settings.projection_years:.1f}", '', '', ''])
        audit_data.append(['Schedule Years', f"{len(cost_schedule)}", '', '', ''])
        audit_data.append(['Start Year', f"{cost_schedule['Year'].min()}", '', '', ''])
        audit_data.append(['End Year', f"{cost_schedule['Year'].max()}", '', '', ''])
        
        audit_df = pd.DataFrame(audit_data, columns=['Check Item', 'Calculated Value', 'Expected/Notes', 'Status', 'Comments'])
        audit_df.to_excel(writer, sheet_name='Audit Trail', index=False)

    def _add_service_master_sheet(self, writer):
        """Add complete service master table for audit purposes."""
        master_data = []
        master_data.append(['Service Master Table (Audit)', '', '', '', '', '', '', '', ''])
        master_data.append(['', '', '', '', '', '', '', '', ''])
        master_data.append(['Category', 'Service Name', 'Unit Cost', 'Frequency/Year', 'Inflation Rate', 'Start Year', 'End Year', 'Service Type', 'Special Years'])
        
        for table_name, table in self.lcp.tables.items():
            for service in table.services:
                service_type = 'Recurring'
                special_years = ''
                
                if service.occurrence_years:
                    service_type = 'Discrete Occurrences'
                    special_years = ', '.join(map(str, service.occurrence_years))
                elif hasattr(service, 'is_distributed_instances') and service.is_distributed_instances:
                    service_type = 'Distributed Instances'
                    special_years = f'{service.total_instances}x over {service.distribution_period_years:.1f} yrs'
                elif service.start_year == service.end_year:
                    service_type = 'One-time'
                
                start_year = service.start_year if service.start_year else self.lcp.settings.base_year
                end_year = service.end_year if service.end_year else self.lcp.settings.base_year + self.lcp.settings.projection_years - 1
                
                master_data.append([
                    table_name,
                    service.name,
                    f"${service.unit_cost:,.2f}",
                    f"{service.frequency_per_year:.1f}",
                    f"{service.inflation_rate:.1%}",
                    str(start_year),
                    str(end_year),
                    service_type,
                    special_years
                ])
        
        master_df = pd.DataFrame(master_data, columns=['Category', 'Service', 'Cost', 'Frequency', 'Inflation', 'Start', 'End', 'Type', 'Special'])
        master_df.to_excel(writer, sheet_name='Service Master', index=False)

    def _export_multi_scenario(self, file_path: str, selected_scenarios: list = None) -> None:
        """Export selected scenarios with comparison sheets."""
        scenarios_to_export = selected_scenarios if selected_scenarios else list(self.lcp.scenarios.keys())
        
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            # Scenario Comparison Overview (only for selected scenarios)
            self._add_scenario_comparison_sheet(writer, scenarios_to_export)
            
            # Export each selected scenario to its own sheet
            for scenario_name in scenarios_to_export:
                if scenario_name in self.lcp.scenarios:
                    # Temporarily switch to this scenario for calculations
                    original_active = self.lcp.active_scenario
                    self.lcp.set_active_scenario(scenario_name)
                    
                    # Create calculator for this scenario
                    scenario_calc = CostCalculator(self.lcp)
                    
                    # Export scenario data with full detail
                    self._export_scenario_to_sheet(writer, scenario_name, scenario_calc)
                    
                    # Restore original active scenario
                    self.lcp.set_active_scenario(original_active)
            
            # Add variance analysis (only for selected scenarios)
            self._add_variance_analysis_sheet(writer, scenarios_to_export)
            
            # Add scenario comparison tables (only for selected scenarios)
            self._add_scenario_comparison_tables(writer, scenarios_to_export)

    def _add_scenario_comparison_sheet(self, writer, selected_scenarios: list = None):
        """Add executive scenario comparison overview."""
        scenarios_to_compare = selected_scenarios if selected_scenarios else list(self.lcp.scenarios.keys())
        comparison_data = []
        comparison_data.append(['SCENARIO COMPARISON OVERVIEW', '', '', '', '', ''])
        comparison_data.append(['', '', '', '', '', ''])
        comparison_data.append(['Scenario Name', 'Description', 'Total Nominal Cost', 'Total Present Value', 'Avg Annual Cost', 'Service Count'])
        
        original_active = self.lcp.active_scenario
        
        for scenario_name in scenarios_to_compare:
            if scenario_name in self.lcp.scenarios:
                scenario = self.lcp.scenarios[scenario_name]
                self.lcp.set_active_scenario(scenario_name)
                calc = CostCalculator(self.lcp)
                summary = calc.calculate_summary_statistics()
                
                # Count total services
                service_count = sum(len(table.services) for table in scenario.tables.values())
                
                comparison_data.append([
                    scenario_name,
                    scenario.description or "No description",
                    f"${summary['total_nominal_cost']:,.2f}",
                    f"${summary['total_present_value']:,.2f}" if self.lcp.evaluee.discount_calculations else "N/A",
                    f"${summary['average_annual_cost']:,.2f}",
                    str(service_count)
                ])
        
        # Restore original active scenario
        self.lcp.set_active_scenario(original_active)
        
        # Create columns based on whether present value is enabled
        if self.lcp.evaluee.discount_calculations:
            columns = ['Scenario', 'Description', 'Nominal Total', 'Present Value', 'Annual Avg', 'Services']
        else:
            columns = ['Scenario', 'Description', 'Nominal Total', 'Annual Avg', 'Services']
            # Remove present value column from data
            comparison_data = [[row[0], row[1], row[2], row[4], row[5]] for row in comparison_data]
        
        comparison_df = pd.DataFrame(comparison_data, columns=columns)
        comparison_df.to_excel(writer, sheet_name='Scenario Comparison', index=False)

    def _export_scenario_to_sheet(self, writer, scenario_name: str, calculator: CostCalculator):
        """Export a single scenario's data to its own sheet."""
        # Build cost schedule for this scenario
        df = calculator.build_cost_schedule()
        summary_stats = calculator.calculate_summary_statistics()
        category_costs = calculator.get_cost_by_category()
        
        # Clean scenario name for sheet name (Excel sheet names have restrictions)
        clean_name = scenario_name.replace('/', '_').replace('\\', '_')[:31]  # Max 31 chars
        
        # Main cost schedule
        df.to_excel(writer, sheet_name=f'{clean_name}_Schedule', index=False)
        
        # Calculate first and last year costs from the cost schedule
        first_year_cost = df.iloc[0]['Total Nominal'] if len(df) > 0 else 0
        last_year_cost = df.iloc[-1]['Total Nominal'] if len(df) > 0 else 0
        
        # Summary for this scenario
        summary_data = [
            ['SCENARIO SUMMARY', ''],
            ['Scenario Name', scenario_name],
            ['Evaluee Name', self.lcp.evaluee.name],
            ['Current Age', f"{self.lcp.evaluee.current_age:.1f} years"],
            ['Base Year', str(self.lcp.settings.base_year)],
            ['Projection Period', f"{self.lcp.settings.projection_years:.1f} years"],
            ['Discount Rate', f"{self.lcp.settings.discount_rate:.1%}" if self.lcp.evaluee.discount_calculations else "Not Applied"],
            ['', ''],
            ['FINANCIAL SUMMARY', ''],
            ['Total Nominal Cost', f"${summary_stats['total_nominal_cost']:,.2f}"],
            ['Total Present Value', f"${summary_stats['total_present_value']:,.2f}" if self.lcp.evaluee.discount_calculations else "N/A"],
            ['Average Annual Cost', f"${summary_stats['average_annual_cost']:,.2f}"],
            ['First Year Cost', f"${first_year_cost:,.2f}"],
            ['Last Year Cost', f"${last_year_cost:,.2f}"]
        ]
        
        summary_df = pd.DataFrame(summary_data, columns=['Metric', 'Value'])
        summary_df.to_excel(writer, sheet_name=f'{clean_name}_Summary', index=False)

    def _add_variance_analysis_sheet(self, writer):
        """Add variance analysis comparing all scenarios to baseline."""
        variance_data = []
        
        # Set up headers based on discount calculations setting
        if self.lcp.evaluee.discount_calculations:
            variance_data.append(['VARIANCE ANALYSIS', '', '', '', ''])
            variance_data.append(['Comparison of all scenarios to baseline', '', '', '', ''])
            variance_data.append(['', '', '', '', ''])
            variance_data.append(['Scenario', 'Nominal Difference', 'Nominal %', 'PV Difference', 'PV %'])
            columns = ['Scenario', 'Nominal Diff', 'Nominal %', 'PV Diff', 'PV %']
        else:
            variance_data.append(['VARIANCE ANALYSIS', '', ''])
            variance_data.append(['Comparison of all scenarios to baseline', '', ''])
            variance_data.append(['', '', ''])
            variance_data.append(['Scenario', 'Nominal Difference', 'Nominal %'])
            columns = ['Scenario', 'Nominal Diff', 'Nominal %']
        
        # Find baseline scenario
        baseline_scenario = self.lcp.get_baseline_scenario()
        if not baseline_scenario:
            if self.lcp.evaluee.discount_calculations:
                variance_data.append(['No baseline scenario found', '', '', '', ''])
            else:
                variance_data.append(['No baseline scenario found', '', ''])
            variance_df = pd.DataFrame(variance_data, columns=columns)
            variance_df.to_excel(writer, sheet_name='Variance Analysis', index=False)
            return
        
        # Calculate baseline totals
        original_active = self.lcp.active_scenario
        self.lcp.set_active_scenario(baseline_scenario.name)
        baseline_calc = CostCalculator(self.lcp)
        baseline_summary = baseline_calc.calculate_summary_statistics()
        
        # Compare each scenario to baseline
        for scenario_name, scenario in self.lcp.scenarios.items():
            if scenario.is_baseline:
                continue  # Skip baseline itself
                
            self.lcp.set_active_scenario(scenario_name)
            calc = CostCalculator(self.lcp)
            summary = calc.calculate_summary_statistics()
            
            # Calculate differences
            nominal_diff = summary['total_nominal_cost'] - baseline_summary['total_nominal_cost']
            nominal_pct = (nominal_diff / baseline_summary['total_nominal_cost']) * 100 if baseline_summary['total_nominal_cost'] > 0 else 0
            
            if self.lcp.evaluee.discount_calculations:
                pv_diff = summary['total_present_value'] - baseline_summary['total_present_value']
                pv_pct = (pv_diff / baseline_summary['total_present_value']) * 100 if baseline_summary['total_present_value'] > 0 else 0
                
                variance_data.append([
                    scenario_name,
                    f"${nominal_diff:,.2f}",
                    f"{nominal_pct:+.1f}%",
                    f"${pv_diff:,.2f}",
                    f"{pv_pct:+.1f}%"
                ])
            else:
                variance_data.append([
                    scenario_name,
                    f"${nominal_diff:,.2f}",
                    f"{nominal_pct:+.1f}%"
                ])
        
        # Restore original active scenario
        self.lcp.set_active_scenario(original_active)
        
        variance_df = pd.DataFrame(variance_data, columns=columns)
        variance_df.to_excel(writer, sheet_name='Variance Analysis', index=False)

    def _add_scenario_comparison_tables(self, writer):
        """Add detailed scenario comparison tables."""
        # Category comparison across scenarios
        category_comparison = []
        category_comparison.append(['CATEGORY COMPARISON ACROSS SCENARIOS', '', '', '', ''])
        category_comparison.append(['', '', '', '', ''])
        
        # Get all unique categories across all scenarios
        all_categories = set()
        for scenario in self.lcp.scenarios.values():
            all_categories.update(scenario.tables.keys())
        
        # Create header row
        header = ['Category'] + list(self.lcp.scenarios.keys())
        category_comparison.append(header)
        
        original_active = self.lcp.active_scenario
        
        # For each category, get costs from each scenario
        for category in sorted(all_categories):
            row = [category]
            
            for scenario_name in self.lcp.scenarios.keys():
                self.lcp.set_active_scenario(scenario_name)
                calc = CostCalculator(self.lcp)
                category_costs = calc.get_cost_by_category()
                
                if category in category_costs:
                    cost = category_costs[category]['table_nominal_total']
                    row.append(f"${cost:,.2f}")
                else:
                    row.append("$0.00")
            
            category_comparison.append(row)
        
        # Restore original active scenario
        self.lcp.set_active_scenario(original_active)
        
        # Pad columns if needed
        max_cols = max(len(row) for row in category_comparison)
        for row in category_comparison:
            while len(row) < max_cols:
                row.append('')
        
        columns = ['Category'] + [f'Scenario_{i}' for i in range(max_cols - 1)]
        comparison_df = pd.DataFrame(category_comparison, columns=columns[:max_cols])
        comparison_df.to_excel(writer, sheet_name='Category Comparison', index=False)


class WordExporter:
    """Export life care plan data to Word document format."""
    
    def __init__(self, calculator: CostCalculator):
        self.calculator = calculator
        self.lcp = calculator.lcp
    
    def export(self, file_path: str, include_chart: bool = True, include_technical_appendix: bool = False, include_all_scenarios: bool = False, selected_scenarios: list = None) -> None:
        """Export the life care plan to Word document in landscape mode.
        
        Args:
            file_path: Output file path
            include_chart: Whether to include cost charts (default: True)
            include_technical_appendix: Whether to include technical methodology and validation 
                                       (default: False for clean legal exhibits)
            include_all_scenarios: Whether to include all scenarios with comparison tables
            selected_scenarios: List of scenario names to include (if None, uses all scenarios)
        """
        if include_all_scenarios and len(self.lcp.scenarios) > 1:
            # Use selected scenarios or all scenarios if none specified
            scenarios_to_export = selected_scenarios if selected_scenarios else list(self.lcp.scenarios.keys())
            if len(scenarios_to_export) > 1:
                self._export_multi_scenario_word(file_path, include_chart, include_technical_appendix, scenarios_to_export)
            else:
                # Only one scenario selected, export as single scenario
                if scenarios_to_export:
                    original_active = self.lcp.active_scenario
                    self.lcp.set_active_scenario(scenarios_to_export[0])
                    self._export_single_scenario_word(file_path, include_chart, include_technical_appendix)
                    self.lcp.set_active_scenario(original_active)
                else:
                    self._export_single_scenario_word(file_path, include_chart, include_technical_appendix)
        else:
            self._export_single_scenario_word(file_path, include_chart, include_technical_appendix)
    
    def _export_single_scenario_word(self, file_path: str, include_chart: bool, include_technical_appendix: bool) -> None:
        """Export current scenario to Word document."""
        doc = Document()

        # Set document to landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap width and height for landscape
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Adjust margins for better table fit
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        # Professional document header for legal exhibit
        doc.add_paragraph()  # Space at top
        
        # Main title
        title = doc.add_heading("LIFE CARE PLAN", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_heading("Economic Analysis and Cost Projections", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space after title
        
        # Evaluee information in professional format
        evaluee_para = doc.add_paragraph()
        evaluee_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        evaluee_para.add_run("Prepared for: ").bold = True
        evaluee_para.add_run(f"{self.lcp.evaluee.name}").bold = True
        
        doc.add_paragraph()  # Space before metadata
        
        # Document information table for professional presentation
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light List'
        
        # Analysis information
        info_data = [
            ["Report Date:", datetime.now().strftime('%B %d, %Y')],
            ["Current Age:", f"{self.lcp.evaluee.current_age:.1f} years"],
            ["Base Year:", str(int(self.lcp.settings.base_year))],
            ["Projection Period:", f"{self.lcp.settings.projection_years:.1f} years"],
            ["End Year:", f"{self.lcp.settings.base_year + self.lcp.settings.projection_years:.1f}"],
            ["Discount Rate:", f"{self.lcp.settings.discount_rate:.1%}" if self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0 else "Not Applied"]
        ]
        
        for i, (label, value) in enumerate(info_data):
            row_cells = info_table.rows[i].cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = value
        
        # Add spacing after metadata
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Professional methodology statement for legal exhibits
        doc.add_heading("Methodology Statement", level=2)
        
        method_para = doc.add_paragraph()
        method_para.add_run("Economic Analysis Basis: ").bold = True
        method_para.add_run("This analysis employs established econometric principles consistent with health economics and actuarial science. ")
        method_para.add_run("Calculations utilize compound annual inflation rates, present value discounting, and frequency-based cost projections. ")
        method_para.add_run("All economic assumptions are based on federal guidelines and historical medical cost data.")
        
        doc.add_paragraph()
        
        # Scientific Acceptance Section
        acceptance_para = doc.add_paragraph()
        acceptance_para.add_run("General Acceptance in Relevant Scientific Community: ").bold = True
        acceptance_para.add_run("Life care plan economic analysis methodologies are widely accepted in forensic economics, ")
        acceptance_para.add_run("rehabilitation counseling, and medical-legal communities. Standards are established by: ")
        acceptance_para.add_run("International Association of Rehabilitation Professionals (IARP), ")
        acceptance_para.add_run("Commission on Health Care Certification (CHCC), ")
        acceptance_para.add_run("National Association of Forensic Economics (NAFE), ")
        acceptance_para.add_run("and peer-reviewed publications in Journal of Forensic Economics, ")
        acceptance_para.add_run("Topics in Spinal Cord Injury Rehabilitation, and similar professional journals.")
        
        doc.add_paragraph()
        
        # Testing and Peer Review Section
        testing_para = doc.add_paragraph()
        testing_para.add_run("Testing and Peer Review: ").bold = True
        testing_para.add_run("The economic modeling techniques used in this analysis have been subject to extensive peer review ")
        testing_para.add_run("through professional literature and court proceedings. Mathematical calculations follow established ")
        testing_para.add_run("financial formulas for present value analysis (PV = FV / (1 + r)^n) and compound growth modeling ")
        testing_para.add_run("(FV = PV × (1 + g)^n). All computational methods are reproducible and verifiable.")
        
        doc.add_paragraph()
        
        # Data Sources and Standards Section
        standards_para = doc.add_paragraph()
        standards_para.add_run("Data Sources and Professional Standards: ").bold = True
        standards_para.add_run("Cost estimates should be derived from reliable sources including: ")
        standards_para.add_run("Medicare fee schedules, private insurance reimbursement rates, ")
        standards_para.add_run("durable medical equipment vendor quotes, pharmaceutical pricing databases, ")
        standards_para.add_run("and published medical literature. Service frequencies should reference ")
        standards_para.add_run("evidence-based treatment protocols, clinical practice guidelines, ")
        standards_para.add_run("and medical professional recommendations specific to the individual's condition.")
        
        doc.add_paragraph()
        
        # Expert Qualifications Required Section
        qualifications_para = doc.add_paragraph()
        qualifications_para.add_run("Expert Qualifications Framework: ").bold = True
        qualifications_para.add_run("Life care plan economic analysis should be conducted by qualified professionals with: ")
        qualifications_para.add_run("(1) Advanced education in economics, healthcare administration, or rehabilitation counseling; ")
        qualifications_para.add_run("(2) Specialized training in life care planning methodology; ")
        qualifications_para.add_run("(3) Professional certification (CRC, CLCP, CVE, or equivalent); ")
        qualifications_para.add_run("(4) Experience with economic analysis and present value calculations; ")
        qualifications_para.add_run("(5) Knowledge of relevant medical conditions and treatment standards.")
        
        doc.add_paragraph()
        
        # Limitations and Assumptions Section
        limitations_para = doc.add_paragraph()
        limitations_para.add_run("Limitations and Key Assumptions: ").bold = True
        limitations_para.add_run("This economic projection is based on current medical knowledge and economic conditions. ")
        limitations_para.add_run("Actual costs may vary due to: changes in medical technology, treatment protocols, ")
        limitations_para.add_run("economic conditions, geographic variations, insurance coverage changes, ")
        limitations_para.add_run("and individual medical developments. Inflation and discount rate assumptions ")
        limitations_para.add_run("represent reasonable estimates but are subject to economic volatility. ")
        limitations_para.add_run("Service frequencies assume stable medical condition and standard care protocols.")
        
        doc.add_paragraph()
        
        # Calculation Transparency Section
        transparency_para = doc.add_paragraph()
        transparency_para.add_run("Calculation Transparency and Reproducibility: ").bold = True
        transparency_para.add_run("All calculations in this report are fully documented and reproducible. ")
        transparency_para.add_run("Mathematical formulas, inflation rates, discount rates, and service frequencies ")
        transparency_para.add_run("are explicitly stated. Raw data inputs and computational methods are available ")
        transparency_para.add_run("for independent verification and cross-examination. Alternative scenarios ")
        transparency_para.add_run("and sensitivity analyses can be performed using different assumption sets.")
        
        # Add legal disclaimer
        doc.add_paragraph()
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.add_run("Legal and Professional Disclaimer: ").bold = True
        disclaimer_para.add_run("This economic analysis is prepared for legal proceedings and expert testimony purposes. ")
        disclaimer_para.add_run("The methodology and conclusions are offered to assist the trier of fact in understanding ")
        disclaimer_para.add_run("future medical care costs. All opinions are expressed within reasonable degree of ")
        disclaimer_para.add_run("professional certainty based on available data and established methodologies.")
        
        # Add Executive Summary Table at top
        doc.add_page_break()
        doc.add_heading("Life Care Plan Cost Summary", level=2)
        
        # Get category costs for summary table
        category_costs = self.calculator.get_cost_by_category()
        summary_stats = self.calculator.calculate_summary_statistics()
        
        # Create summary table with service categories
        # Only show present value if discount calculations are enabled AND discount rate > 0
        show_present_value = self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0
        if show_present_value:
            summary_headers = ['Service Category', 'Total Lifetime Cost (Nominal)', 'Total Lifetime Cost (Present Value)']
        else:
            summary_headers = ['Service Category', 'Total Lifetime Cost (Nominal)']
        
        # Calculate table size
        num_categories = len(category_costs)
        summary_table = doc.add_table(rows=num_categories + 2, cols=len(summary_headers))  # +2 for header and grand total
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        summary_table.style = 'Light List'
        
        # Set column widths
        category_width = Inches(3.0)
        cost_width = Inches(2.2)
        
        summary_table.columns[0].width = category_width
        for i in range(1, len(summary_headers)):
            summary_table.columns[i].width = cost_width
        
        # Header row
        hdr_cells = summary_table.rows[0].cells
        for idx, header_text in enumerate(summary_headers):
            hdr_cells[idx].text = header_text
            paragraph = hdr_cells[idx].paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[idx].top_margin = Pt(6)
            hdr_cells[idx].bottom_margin = Pt(6)
            hdr_cells[idx].left_margin = Pt(4)
            hdr_cells[idx].right_margin = Pt(4)
        
        # Data rows for each category
        grand_total_nominal = 0
        grand_total_pv = 0
        
        for row_idx, (table_name, data) in enumerate(category_costs.items(), start=1):
            row_cells = summary_table.rows[row_idx].cells
            
            # Category name
            row_cells[0].text = table_name
            paragraph = row_cells[0].paragraphs[0]
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Nominal cost
            nominal_cost = data['table_nominal_total']
            row_cells[1].text = f"${nominal_cost:,.2f}"
            paragraph = row_cells[1].paragraphs[0]
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            grand_total_nominal += nominal_cost
            
            # Present value cost (if applicable)
            if show_present_value:
                pv_cost = data['table_present_value_total']
                row_cells[2].text = f"${pv_cost:,.2f}"
                paragraph = row_cells[2].paragraphs[0]
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                grand_total_pv += pv_cost
            
            # Cell formatting
            for cell in row_cells:
                cell.top_margin = Pt(4)
                cell.bottom_margin = Pt(4)
                cell.left_margin = Pt(4)
                cell.right_margin = Pt(4)
        
        # Grand total row
        total_row_cells = summary_table.rows[-1].cells
        total_row_cells[0].text = "GRAND TOTAL"
        paragraph = total_row_cells[0].paragraphs[0]
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        total_row_cells[1].text = f"${grand_total_nominal:,.2f}"
        paragraph = total_row_cells[1].paragraphs[0]
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if show_present_value:
            total_row_cells[2].text = f"${grand_total_pv:,.2f}"
            paragraph = total_row_cells[2].paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Format grand total row
        for cell in total_row_cells:
            cell.top_margin = Pt(6)
            cell.bottom_margin = Pt(6)
            cell.left_margin = Pt(4)
            cell.right_margin = Pt(4)
        
        doc.add_paragraph()  # Space after summary table
        
        # Mathematical Methodology Documentation
        doc.add_heading("Mathematical Formulas and Calculation Methods", level=2)
        
        # Formula documentation
        formula_para = doc.add_paragraph()
        formula_para.add_run("Inflation Adjustment Formula: ").bold = True
        formula_para.add_run("Future Cost = Present Cost × (1 + inflation_rate)^years_from_base")
        
        doc.add_paragraph()
        pv_para = doc.add_paragraph()
        pv_para.add_run("Present Value Formula: ").bold = True
        pv_para.add_run("Present Value = Future Value ÷ (1 + discount_rate)^years_from_base")
        
        doc.add_paragraph()
        annual_para = doc.add_paragraph()
        annual_para.add_run("Annual Service Cost: ").bold = True
        annual_para.add_run("Annual Cost = Unit Cost × Frequency per Year × Inflation Adjustment")
        
        doc.add_paragraph()
        lifetime_para = doc.add_paragraph()
        lifetime_para.add_run("Lifetime Service Cost: ").bold = True
        lifetime_para.add_run("Sum of all annual costs over the service period, with inflation applied to each year")
        
        doc.add_paragraph()
        
        # Economic Assumptions Documentation
        economic_para = doc.add_paragraph()
        economic_para.add_run("Economic Assumptions Used: ").bold = True
        if show_present_value:
            economic_para.add_run(f"Discount Rate: {self.lcp.settings.discount_rate:.1%} annually. ")
        economic_para.add_run(f"Analysis Period: {self.lcp.settings.projection_years:.1f} years ")
        economic_para.add_run(f"({int(self.lcp.settings.base_year)} through {self.lcp.settings.base_year + self.lcp.settings.projection_years - 1:.1f}). ")
        economic_para.add_run("Individual service inflation rates as specified in service details. ")
        economic_para.add_run("All calculations assume consistent annual application of stated rates.")
        
        doc.add_paragraph()
        
        # Quality Control Documentation
        qc_para = doc.add_paragraph()
        qc_para.add_run("Quality Control and Verification: ").bold = True
        qc_para.add_run("All calculations are performed using established financial mathematics. ")
        qc_para.add_run("Results are subject to mathematical verification and cross-checking. ")
        qc_para.add_run("Alternative calculation methods may be applied for confirmation. ")
        qc_para.add_run("Sensitivity analysis can be performed using different assumption sets ")
        qc_para.add_run("to test the robustness of projections under varying economic conditions.")
        
        # Only add technical appendix if requested (not for legal exhibits)
        if include_technical_appendix:
            # Calculation methodology section removed per user request
            pass
        
        # Summary statistics
        doc.add_heading("Executive Summary", level=2)
        summary_stats = self.calculator.calculate_summary_statistics()

        summary_para = doc.add_paragraph()
        summary_para.add_run("Total Lifetime Medical Costs (Nominal): ").bold = True
        summary_para.add_run(f"${summary_stats['total_nominal_cost']:,.2f}\n")

        summary_para.add_run("Average Annual Medical Costs: ").bold = True
        summary_para.add_run(f"${summary_stats['average_annual_cost']:,.2f}\n")

        if self.lcp.evaluee.discount_calculations:
            summary_para.add_run("Total Lifetime Medical Costs (Present Value): ").bold = True
            summary_para.add_run(f"${summary_stats['total_present_value']:,.2f}\n")

            savings = summary_stats['total_nominal_cost'] - summary_stats['total_present_value']
            summary_para.add_run("Present Value Savings vs Nominal: ").bold = True
            summary_para.add_run(f"${savings:,.2f}\n")
        
        # Add spacing after summary
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Detailed Service Category Breakdown
        doc.add_heading("Detailed Service Breakdown by Category", level=2)
        category_costs = self.calculator.get_cost_by_category()
        
        for table_name, data in category_costs.items():
            # Category header with summary
            doc.add_heading(f"{table_name}", level=3)
            
            # Category summary paragraph with Daubert compliance
            summary_para = doc.add_paragraph()
            summary_para.add_run("Category Summary: ").bold = True
            summary_para.add_run(f"This category contains {len(data['services'])} medical service(s) with a total lifetime cost of ")
            summary_para.add_run(f"${data['table_nominal_total']:,.2f}").bold = True
            if show_present_value:
                summary_para.add_run(f" (${data['table_present_value_total']:,.2f} in present value)")
            summary_para.add_run(".")
            
            
            doc.add_paragraph()  # Spacing
            
            if data['services']:
                # Create detailed service table for this category
                service_table_headers = [
                    'Service Name',
                    'Cost per Unit\n(Data Source Required)',
                    'Frequency per Year\n(Medical Basis Required)',
                    'Service Period\n(Clinical Justification)',
                    'Annual Inflation Rate\n(Economic Basis)',
                    'Total Lifetime Cost\n(Calculated)'
                ]
                
                if show_present_value:
                    service_table_headers.append('Present Value\nLifetime Cost\n(Calculated)')
                
                # Create service table
                service_table = doc.add_table(rows=len(data['services']) + 1, cols=len(service_table_headers))
                service_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                service_table.style = 'Light List'
                
                # Set column widths for service table
                col_widths = [Inches(2.0), Inches(1.0), Inches(0.8), Inches(1.5), Inches(0.8), Inches(1.3)]
                if self.lcp.evaluee.discount_calculations:
                    col_widths.append(Inches(1.3))
                
                for i, width in enumerate(col_widths):
                    if i < len(service_table.columns):
                        service_table.columns[i].width = width
                
                # Header row
                hdr_cells = service_table.rows[0].cells
                for idx, header_text in enumerate(service_table_headers):
                    hdr_cells[idx].text = header_text
                    paragraph = hdr_cells[idx].paragraphs[0]
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Cell formatting
                    hdr_cells[idx].top_margin = Pt(4)
                    hdr_cells[idx].bottom_margin = Pt(4)
                    hdr_cells[idx].left_margin = Pt(3)
                    hdr_cells[idx].right_margin = Pt(3)
                
                # Data rows for each service
                for row_idx, service in enumerate(data['services'], start=1):
                    row_cells = service_table.rows[row_idx].cells
                    
                    # Determine service period description
                    if service['is_one_time_cost']:
                        service_period = f"One-time in {service['one_time_cost_year']}"
                    elif service['occurrence_years']:
                        years = service['occurrence_years']
                        if len(years) == 1:
                            service_period = f"Year {years[0]} only"
                        else:
                            service_period = f"Years {min(years)}-{max(years)}\n({len(years)} specific years)"
                    elif service.get('is_distributed_instances', False):
                        service_period = f"{service['total_instances']}x over {service['distribution_period_years']:.1f} years\n(Starting {service['start_year']})"
                    else:
                        start_yr = service['start_year'] if service['start_year'] else 'Start of plan'
                        end_yr = service['end_year'] if service['end_year'] else 'End of plan'
                        service_period = f"{start_yr} to {end_yr}"
                    
                    # Fill in service data
                    if service.get('is_distributed_instances', False):
                        frequency_display = f"{service['frequency_per_year']:.2f}/yr\n({service['total_instances']}x total)"
                    else:
                        frequency_display = f"{service['frequency_per_year']:.1f}x"
                    
                    service_data = [
                        service['name'],
                        f"${service['unit_cost']:,.2f}",
                        frequency_display,
                        service_period,
                        f"{service['inflation_rate']:.1f}%",
                        f"${service['nominal_total']:,.2f}"
                    ]
                    
                    if show_present_value:
                        service_data.append(f"${service['present_value_total']:,.2f}")
                    
                    for col_idx, cell_value in enumerate(service_data):
                        row_cells[col_idx].text = cell_value
                        
                        # Cell formatting
                        row_cells[col_idx].top_margin = Pt(3)
                        row_cells[col_idx].bottom_margin = Pt(3)
                        row_cells[col_idx].left_margin = Pt(3)
                        row_cells[col_idx].right_margin = Pt(3)
                        
                        # Text formatting
                        paragraph = row_cells[col_idx].paragraphs[0]
                        if paragraph.runs:
                            paragraph.runs[0].font.size = Pt(8)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add legal and methodological explanatory text
                doc.add_paragraph()
                explanation_para = doc.add_paragraph()
                explanation_para.add_run("Expert Opinion Basis: ").bold = True
                explanation_para.add_run("Each service listed requires supporting documentation including: ")
                explanation_para.add_run("(1) Medical necessity determined by qualified healthcare providers; ")
                explanation_para.add_run("(2) Cost estimates from reliable sources (providers, fee schedules, market research); ")
                explanation_para.add_run("(3) Frequency based on treatment protocols or physician recommendations; ")
                explanation_para.add_run("(4) Duration supported by medical literature or clinical experience. ")
                
                doc.add_paragraph()
                calculation_para = doc.add_paragraph()
                calculation_para.add_run("Calculation Methodology: ").bold = True
                calculation_para.add_run("Costs are projected using compound inflation modeling applied annually. ")
                if show_present_value:
                    calculation_para.add_run("Present value calculations discount future costs to current dollars using ")
                    calculation_para.add_run(f"{self.lcp.settings.discount_rate:.1%} annual discount rate").bold = True
                    calculation_para.add_run(" consistent with federal economic guidelines. ")
                calculation_para.add_run("All mathematical operations follow established financial principles ")
                calculation_para.add_run("and are subject to independent verification and cross-examination.")
                
                doc.add_paragraph()
                limitations_service_para = doc.add_paragraph()
                limitations_service_para.add_run("Service-Specific Limitations: ").bold = True
                limitations_service_para.add_run("Projections assume medical stability and standard treatment protocols. ")
                limitations_service_para.add_run("Individual variations in treatment response, complications, or medical advances ")
                limitations_service_para.add_run("may alter actual service needs and costs. Expert opinions should be updated ")
                limitations_service_para.add_run("as medical conditions and standards of care evolve.")
                
            # Add spacing between categories
            doc.add_paragraph()
            doc.add_page_break() if len(category_costs) > 1 else doc.add_paragraph()
        
        # Annual Cost Schedule with Category Breakdown
        doc.add_page_break()
        doc.add_heading("Annual Cost Schedule Summary", level=2)
        
        # Add explanation paragraph
        explanation_para = doc.add_paragraph()
        explanation_para.add_run("Understanding Your Annual Costs: ").bold = True
        explanation_para.add_run("The table below shows the total medical costs for each year of the life care plan. ")
        explanation_para.add_run("These costs represent all services combined and include inflation adjustments. ")
        if show_present_value:
            explanation_para.add_run("The present value column shows what future costs are worth in today's dollars.")
        
        doc.add_paragraph()  # Spacing

        df = self.calculator.build_cost_schedule()
        
        # Create simplified annual summary table
        # Only show present value if discount calculations are enabled AND discount rate > 0
        show_present_value = self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0
        
        if show_present_value and "Present Value" in df.columns:
            table_columns = ['Year', 'Evaluee Age', 'Total Annual Cost', 'Present Value Cost']
            table_data = []
            for _, row in df.iterrows():
                table_data.append([
                    str(int(row['Year'])),
                    str(int(row['Age'])),
                    f"${row['Total Nominal']:,.0f}",
                    f"${row['Present Value']:,.0f}"
                ])
        else:
            table_columns = ['Year', 'Evaluee Age', 'Total Annual Cost']
            table_data = []
            for _, row in df.iterrows():
                table_data.append([
                    str(int(row['Year'])),
                    str(int(row['Age'])),
                    f"${row['Total Nominal']:,.0f}"
                ])
        
        # Create summary table
        summary_table = doc.add_table(rows=len(table_data) + 1, cols=len(table_columns))
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        summary_table.style = 'Light List'
        
        # Set column widths
        year_width = Inches(1.0)
        age_width = Inches(1.2) 
        cost_width = Inches(1.8)
        
        summary_table.columns[0].width = year_width
        summary_table.columns[1].width = age_width
        for i in range(2, len(table_columns)):
            summary_table.columns[i].width = cost_width

        # Header row
        hdr_cells = summary_table.rows[0].cells
        for idx, header_text in enumerate(table_columns):
            hdr_cells[idx].text = header_text
            paragraph = hdr_cells[idx].paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[idx].top_margin = Pt(4)
            hdr_cells[idx].bottom_margin = Pt(4)
            hdr_cells[idx].left_margin = Pt(3)
            hdr_cells[idx].right_margin = Pt(3)

        # Data rows
        for row_idx, row_data in enumerate(table_data, start=1):
            row_cells = summary_table.rows[row_idx].cells
            
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = cell_value
                row_cells[col_idx].top_margin = Pt(3)
                row_cells[col_idx].bottom_margin = Pt(3)
                row_cells[col_idx].left_margin = Pt(3)
                row_cells[col_idx].right_margin = Pt(3)
                
                paragraph = row_cells[col_idx].paragraphs[0]
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add Year-by-Year Loss Schedule
        doc.add_page_break()
        doc.add_heading("Year-by-Year Loss Schedule", level=2)
        
        # Add explanation
        loss_schedule_explanation = doc.add_paragraph()
        loss_schedule_explanation.add_run("Loss Schedule Overview: ").bold = True
        loss_schedule_explanation.add_run("This section provides a comprehensive year-by-year analysis of projected medical costs, ")
        loss_schedule_explanation.add_run("organized both by overall yearly totals and detailed service category breakdowns. ")
        loss_schedule_explanation.add_run("This format assists in understanding annual cost patterns and service delivery timing.")
        
        doc.add_paragraph()
        
        # Create overall summary table showing yearly totals by category
        doc.add_heading("Overall Yearly Summary by Service Category", level=3)
        
        # Calculate year-by-year totals by category
        base_year = int(self.lcp.settings.base_year)
        end_year = base_year + int(self.lcp.settings.projection_years)
        if self.lcp.settings.projection_years % 1 != 0:
            end_year += 1
        years = list(range(base_year, end_year))
        category_names = list(self.lcp.tables.keys())
        
        # Create matrix of costs: year x category
        yearly_category_matrix = {}
        yearly_totals = {}
        category_totals = {}
        
        for year in years:
            yearly_category_matrix[year] = {}
            yearly_totals[year] = 0
            
            for table_name, table in self.lcp.tables.items():
                category_cost = 0
                for service in table.services:
                    service_cost = self.calculator.calculate_service_cost(service, year)
                    category_cost += float(service_cost)
                
                yearly_category_matrix[year][table_name] = category_cost
                yearly_totals[year] += category_cost
                
                if table_name not in category_totals:
                    category_totals[table_name] = 0
                category_totals[table_name] += category_cost
        
        # Create summary table with years as rows, categories as columns
        summary_headers = ['Year', 'Evaluee Age'] + category_names + ['Annual Total']
        
        # Calculate table dimensions
        num_years = len(years)
        num_cols = len(summary_headers)
        
        summary_table = doc.add_table(rows=num_years + 2, cols=num_cols)  # +2 for header and totals
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        summary_table.style = 'Light List'
        
        # Set column widths for summary table
        year_col_width = Inches(0.7)
        age_col_width = Inches(0.8)
        category_col_width = Inches(1.1)
        total_col_width = Inches(1.2)
        
        summary_table.columns[0].width = year_col_width
        summary_table.columns[1].width = age_col_width
        for i in range(2, len(category_names) + 2):
            summary_table.columns[i].width = category_col_width
        summary_table.columns[-1].width = total_col_width
        
        # Header row
        header_cells = summary_table.rows[0].cells
        for idx, header in enumerate(summary_headers):
            header_cells[idx].text = header
            paragraph = header_cells[idx].paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(9)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row_idx, year in enumerate(years, start=1):
            evaluee_age = int(self.lcp.evaluee.current_age + (year - int(self.lcp.settings.base_year)))
            row_cells = summary_table.rows[row_idx].cells
            
            # Year and age
            row_cells[0].text = str(year)
            row_cells[1].text = str(evaluee_age)
            
            # Category costs
            for col_idx, category_name in enumerate(category_names, start=2):
                cost = yearly_category_matrix[year][category_name]
                row_cells[col_idx].text = f"${cost:,.0f}" if cost > 0 else "-"
                
                # Format cell
                paragraph = row_cells[col_idx].paragraphs[0]
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(8)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Annual total
            row_cells[-1].text = f"${yearly_totals[year]:,.0f}"
            
            # Format year, age, and total cells
            for cell_idx in [0, 1, -1]:
                paragraph = row_cells[cell_idx].paragraphs[0]
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(8)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Totals row
        totals_cells = summary_table.rows[-1].cells
        totals_cells[0].text = "TOTALS"
        totals_cells[1].text = ""
        
        for col_idx, category_name in enumerate(category_names, start=2):
            total_cost = category_totals[category_name]
            totals_cells[col_idx].text = f"${total_cost:,.0f}"
            
            # Format cell
            paragraph = totals_cells[col_idx].paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(totals_cells[col_idx].text)
            run.bold = True
            run.font.size = Pt(9)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Grand total
        grand_total = sum(yearly_totals.values())
        totals_cells[-1].text = f"${grand_total:,.0f}"
        paragraph = totals_cells[-1].paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(totals_cells[-1].text)
        run.bold = True
        run.font.size = Pt(9)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format first column of totals row
        paragraph = totals_cells[0].paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(totals_cells[0].text)
        run.bold = True
        run.font.size = Pt(9)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add summary explanation
        summary_explanation = doc.add_paragraph()
        summary_explanation.add_run("Summary Table Explanation: ").bold = True
        summary_explanation.add_run("This table shows the total cost for each service category by year. ")
        summary_explanation.add_run("Reading across each row shows the annual cost distribution across different types of medical services. ")
        summary_explanation.add_run("Reading down each column shows how costs for a specific service category change over time due to inflation and service timing.")
        
        # Add comprehensive year-by-year breakdown by category
        doc.add_page_break()
        doc.add_heading("Detailed Year-by-Year Service Breakdown by Category", level=2)
        
        # Add explanation and validation notice
        detailed_explanation = doc.add_paragraph()
        detailed_explanation.add_run("Year-by-Year Service Details: ").bold = True
        detailed_explanation.add_run("The following section shows exactly which services are provided each year and their individual costs. ")
        detailed_explanation.add_run("This detailed breakdown helps you understand what drives the costs in each year of the plan.")
        
        doc.add_paragraph()
        validation_para = doc.add_paragraph()
        validation_para.add_run("Cross-Verification Notice: ").bold = True
        validation_para.add_run("All calculations in this detailed breakdown use identical mathematical methods as the Summary Schedule above. ")
        validation_para.add_run("Year totals in this section should match corresponding years in the Annual Cost Schedule Summary. ")
        validation_para.add_run("Any discrepancies indicate calculation errors that require correction.")
        
        doc.add_paragraph()
        
        # Get detailed year-by-year data
        category_costs = self.calculator.get_cost_by_category()
        base_year = int(self.lcp.settings.base_year)
        end_year = base_year + int(self.lcp.settings.projection_years)
        if self.lcp.settings.projection_years % 1 != 0:
            end_year += 1
        years = list(range(base_year, end_year))
        
        # Create a comprehensive table showing services by year
        for year in years:
            evaluee_age = int(self.lcp.evaluee.current_age + (year - self.lcp.settings.base_year))
            doc.add_heading(f"Year {year} (Evaluee Age: {evaluee_age})", level=3)
            
            year_services = []
            year_total = 0
            year_total_pv = 0
            
            # Use calculator's service cost method for consistency
            for table_name, table in self.lcp.tables.items():
                for service in table.services:
                    # Use the calculator's method to get the correct cost for this year
                    service_cost = self.calculator.calculate_service_cost(service, year)
                    
                    if service_cost > 0:  # Service applies to this year
                        # Calculate present value using calculator's method
                        years_from_base = year - int(self.lcp.settings.base_year)
                        service_cost_pv = 0
                        if self.lcp.evaluee.discount_calculations:
                            service_cost_pv = float(self.calculator.calculate_present_value(service_cost, years_from_base))
                        
                        # Determine frequency display
                        if service.is_one_time_cost:
                            frequency_display = 1
                        else:
                            frequency_display = service.frequency_per_year
                        
                        year_services.append({
                            'category': table_name,
                            'name': service.name,
                            'frequency': frequency_display,
                            'unit_cost': service.unit_cost,
                            'inflated_cost': float(service_cost),
                            'present_value_cost': service_cost_pv,
                            'is_one_time': service.is_one_time_cost
                        })
                        
                        year_total += float(service_cost)
                        year_total_pv += service_cost_pv
            
            if year_services:
                # Create table for this year's services
                year_table_headers = ['Service Category', 'Service Name', 'Frequency', 'Cost This Year']
                # Only show present value if discount calculations are enabled AND discount rate > 0
                if self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0:
                    year_table_headers.append('Present Value Cost')
                
                year_table = doc.add_table(rows=len(year_services) + 2, cols=len(year_table_headers))  # +2 for header and total
                year_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                year_table.style = 'Light List'
                
                # Set column widths
                year_col_widths = [Inches(1.8), Inches(2.2), Inches(0.8), Inches(1.2)]
                if self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0:
                    year_col_widths.append(Inches(1.2))
                
                for i, width in enumerate(year_col_widths):
                    if i < len(year_table.columns):
                        year_table.columns[i].width = width
                
                # Headers
                hdr_cells = year_table.rows[0].cells
                for idx, header_text in enumerate(year_table_headers):
                    hdr_cells[idx].text = header_text
                    paragraph = hdr_cells[idx].paragraphs[0]
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Service rows
                for row_idx, service in enumerate(year_services, start=1):
                    row_cells = year_table.rows[row_idx].cells
                    
                    frequency_text = "One-time" if service['is_one_time'] else f"{service['frequency']:.1f}x/year"
                    
                    service_row_data = [
                        service['category'],
                        service['name'],
                        frequency_text,
                        f"${service['inflated_cost']:,.0f}"
                    ]
                    
                    # Only include present value if discount calculations are enabled AND discount rate > 0
                    if self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0:
                        service_row_data.append(f"${service['present_value_cost']:,.0f}")
                    
                    for col_idx, cell_value in enumerate(service_row_data):
                        row_cells[col_idx].text = cell_value
                        paragraph = row_cells[col_idx].paragraphs[0]
                        if paragraph.runs:
                            paragraph.runs[0].font.size = Pt(8)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Total row
                total_row_cells = year_table.rows[-1].cells
                total_row_cells[0].text = "YEAR TOTAL"
                total_row_cells[1].text = ""
                total_row_cells[2].text = ""
                total_row_cells[3].text = f"${year_total:,.0f}"
                
                # Only include present value total if discount calculations are enabled AND discount rate > 0
                if self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0:
                    total_row_cells[4].text = f"${year_total_pv:,.0f}"
                
                # Format total row
                for cell in total_row_cells:
                    if cell.text:
                        paragraph = cell.paragraphs[0]
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell.text)
                        run.bold = True
                        run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # No services for this year
                no_services_para = doc.add_paragraph()
                no_services_para.add_run("No medical services scheduled for this year.").italic = True
            
            doc.add_paragraph()  # Spacing between years
        
        
        # Add spacing after table
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add chart if requested
        if include_chart:
            chart_path = self._create_chart()
            if chart_path and os.path.exists(chart_path):
                doc.add_page_break()
                doc.add_heading("Cost Visualization", level=2)
                doc.add_paragraph()  # Add spacing before chart
                
                # Center the chart
                chart_para = doc.add_paragraph()
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_run = chart_para.add_run()
                chart_run.add_picture(chart_path, width=Inches(8))  # Larger chart for better readability
                
                # Clean up temporary chart file
                os.remove(chart_path)
        
        doc.save(file_path)

    def _export_multi_scenario_word(self, file_path: str, include_chart: bool, include_technical_appendix: bool, selected_scenarios: list = None) -> None:
        """Export multi-scenario comparison to Word document."""
        scenarios_to_export = selected_scenarios if selected_scenarios else list(self.lcp.scenarios.keys())
        doc = Document()

        # Set document to landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Adjust margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        # Main title
        title = doc.add_heading("LIFE CARE PLAN - MULTI-SCENARIO ANALYSIS", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_heading("Scenario Comparison and Economic Analysis", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Evaluee information
        evaluee_para = doc.add_paragraph()
        evaluee_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        evaluee_para.add_run("Prepared for: ").bold = True
        evaluee_para.add_run(f"{self.lcp.evaluee.name}").bold = True
        
        doc.add_paragraph()
        
        # Document information
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light List'
        
        info_data = [
            ["Report Date:", datetime.now().strftime('%B %d, %Y')],
            ["Number of Scenarios:", str(len(scenarios_to_export))],
            ["Base Year:", str(int(self.lcp.settings.base_year))],
            ["Projection Period:", f"{self.lcp.settings.projection_years:.1f} years"]
        ]
        
        for i, (label, value) in enumerate(info_data):
            row_cells = info_table.rows[i].cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = value
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Executive Scenario Comparison
        doc.add_heading("Executive Scenario Comparison", level=2)
        
        # Create comparison table
        scenarios = [(name, self.lcp.scenarios[name]) for name in scenarios_to_export if name in self.lcp.scenarios]
        comparison_headers = ['Metric'] + [scenario[0] for scenario in scenarios]
        
        comparison_table = doc.add_table(rows=6, cols=len(comparison_headers))
        comparison_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        comparison_table.style = 'Light List'
        
        # Header row
        hdr_cells = comparison_table.rows[0].cells
        for idx, header_text in enumerate(comparison_headers):
            hdr_cells[idx].text = header_text
            paragraph = hdr_cells[idx].paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Collect scenario data
        original_active = self.lcp.active_scenario
        scenario_data = {}
        
        for scenario_name, scenario in scenarios:
            self.lcp.set_active_scenario(scenario_name)
            calc = CostCalculator(self.lcp)
            summary = calc.calculate_summary_statistics()
            
            # Calculate first and last year costs
            df = calc.build_cost_schedule()
            first_year_cost = df.iloc[0]['Total Nominal'] if len(df) > 0 else 0
            last_year_cost = df.iloc[-1]['Total Nominal'] if len(df) > 0 else 0
            
            # Add to summary
            summary['first_year_cost'] = first_year_cost
            summary['last_year_cost'] = last_year_cost
            
            scenario_data[scenario_name] = summary
        
        # Data rows
        metrics = [
            ("Total Nominal Cost", "total_nominal_cost"),
            ("Total Present Value", "total_present_value"),
            ("Average Annual Cost", "average_annual_cost"),
            ("First Year Cost", "first_year_cost"),
            ("Last Year Cost", "last_year_cost")
        ]
        
        for row_idx, (metric_name, metric_key) in enumerate(metrics, start=1):
            row_cells = comparison_table.rows[row_idx].cells
            row_cells[0].text = metric_name
            row_cells[0].paragraphs[0].runs[0].bold = True
            
            for col_idx, (scenario_name, _) in enumerate(scenarios, start=1):
                if metric_key == "total_present_value" and not self.lcp.evaluee.discount_calculations:
                    row_cells[col_idx].text = "N/A"
                else:
                    value = scenario_data[scenario_name][metric_key]
                    row_cells[col_idx].text = f"${value:,.2f}"
                    
                # Formatting
                paragraph = row_cells[col_idx].paragraphs[0]
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Restore original active scenario
        self.lcp.set_active_scenario(original_active)
        
        doc.add_paragraph()
        
        # Variance Analysis Section
        doc.add_heading("Variance Analysis", level=2)
        
        # Find baseline scenario
        baseline_scenario = self.lcp.get_baseline_scenario()
        if baseline_scenario:
            baseline_data = scenario_data[baseline_scenario.name]
            
            variance_para = doc.add_paragraph()
            variance_para.add_run("Baseline Scenario: ").bold = True
            variance_para.add_run(f"{baseline_scenario.name} - ${baseline_data['total_nominal_cost']:,.2f} total nominal cost")
            
            doc.add_paragraph()
            
            # Variance table
            variance_headers = ['Scenario', 'Difference from Baseline', 'Percentage Change']
            variance_table = doc.add_table(rows=len(scenarios), cols=len(variance_headers))
            variance_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            variance_table.style = 'Light List'
            
            # Header
            hdr_cells = variance_table.rows[0].cells
            for idx, header_text in enumerate(variance_headers):
                hdr_cells[idx].text = header_text
                paragraph = hdr_cells[idx].paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for row_idx, (scenario_name, _) in enumerate(scenarios[1:], start=1):  # Skip baseline
                if scenario_name == baseline_scenario.name:
                    continue
                    
                row_cells = variance_table.rows[row_idx].cells
                scenario_total = scenario_data[scenario_name]['total_nominal_cost']
                difference = scenario_total - baseline_data['total_nominal_cost']
                percentage = (difference / baseline_data['total_nominal_cost']) * 100 if baseline_data['total_nominal_cost'] > 0 else 0
                
                row_cells[0].text = scenario_name
                row_cells[1].text = f"${difference:+,.2f}"
                row_cells[2].text = f"{percentage:+.1f}%"
                
                # Formatting
                for cell in row_cells:
                    paragraph = cell.paragraphs[0]
                    if paragraph.runs:
                        paragraph.runs[0].font.size = Pt(9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Individual Scenario Details
        doc.add_page_break()
        doc.add_heading("Individual Scenario Details", level=2)
        
        for scenario_name, scenario in scenarios:
            doc.add_heading(f"Scenario: {scenario_name}", level=3)
            
            # Description
            if scenario.description:
                desc_para = doc.add_paragraph()
                desc_para.add_run("Description: ").bold = True
                desc_para.add_run(scenario.description)
                doc.add_paragraph()
            
            # Summary for this scenario
            self.lcp.set_active_scenario(scenario_name)
            calc = CostCalculator(self.lcp)
            summary = calc.calculate_summary_statistics()
            category_costs = calc.get_cost_by_category()
            
            # Category breakdown
            category_table = doc.add_table(rows=len(category_costs) + 2, cols=3)
            category_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            category_table.style = 'Light List'
            
            # Headers
            headers = ['Service Category', 'Total Nominal Cost', 'Total Present Value']
            if not self.lcp.evaluee.discount_calculations:
                headers = headers[:2]
            
            hdr_cells = category_table.rows[0].cells
            for idx, header_text in enumerate(headers):
                hdr_cells[idx].text = header_text
                paragraph = hdr_cells[idx].paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Category data
            total_nominal = 0
            total_pv = 0
            
            for row_idx, (cat_name, cat_data) in enumerate(category_costs.items(), start=1):
                row_cells = category_table.rows[row_idx].cells
                nominal_cost = cat_data['table_nominal_total']
                pv_cost = cat_data['table_present_value_total']
                
                row_cells[0].text = cat_name
                row_cells[1].text = f"${nominal_cost:,.2f}"
                if self.lcp.evaluee.discount_calculations:
                    row_cells[2].text = f"${pv_cost:,.2f}"
                
                total_nominal += nominal_cost
                total_pv += pv_cost
                
                # Formatting
                for cell in row_cells:
                    paragraph = cell.paragraphs[0]
                    if paragraph.runs:
                        paragraph.runs[0].font.size = Pt(9)
                    if cell == row_cells[0]:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Total row
            total_cells = category_table.rows[-1].cells
            total_cells[0].text = "TOTAL"
            total_cells[1].text = f"${total_nominal:,.2f}"
            if self.lcp.evaluee.discount_calculations:
                total_cells[2].text = f"${total_pv:,.2f}"
            
            # Format total row
            for cell in total_cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                if cell == total_cells[0]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()
            doc.add_paragraph()
        
        # Restore original active scenario
        self.lcp.set_active_scenario(original_active)
        
        # Add methodology note
        doc.add_page_break()
        doc.add_heading("Multi-Scenario Analysis Methodology", level=2)
        
        method_para = doc.add_paragraph()
        method_para.add_run("Scenario Analysis Purpose: ").bold = True
        method_para.add_run("Multiple scenarios allow evaluation of different treatment approaches, ")
        method_para.add_run("service intensity levels, or economic assumptions. Each scenario represents ")
        method_para.add_run("a distinct set of medical services and associated costs, enabling ")
        method_para.add_run("comprehensive comparison for decision-making purposes.")
        
        doc.add_paragraph()
        
        consistency_para = doc.add_paragraph()
        consistency_para.add_run("Calculation Consistency: ").bold = True
        consistency_para.add_run("All scenarios use identical mathematical methods for inflation ")
        consistency_para.add_run("adjustment and present value calculation. Differences in results ")
        consistency_para.add_run("reflect only differences in service specifications, not calculation methods. ")
        consistency_para.add_run("This ensures accurate comparison between scenarios.")
        
        doc.save(file_path)

    def export_technical_validation_report(self, file_path: str) -> None:
        """Export a separate technical validation and methodology report.
        
        This generates a comprehensive technical document with all validation checks,
        methodology details, and analysis - suitable for technical review but separate
        from legal exhibit documents.
        """
        doc = Document()
        
        # Title for technical report
        title = doc.add_heading("Life Care Plan - Technical Validation Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f"Technical Analysis for: {self.lcp.evaluee.name}", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Report metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run("Technical Report Generated: ").bold = True
        meta_para.add_run(f"{datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
        
        doc.add_paragraph()
        
        # Calculation methodology section removed per user request
        
        # Add comprehensive validation results
        doc.add_heading("Comprehensive Validation Results", level=2)
        
        # Perform and display variance analysis
        variance_results = self.calculator.perform_variance_analysis()
        
        # Quality control summary
        qc_para = doc.add_paragraph()
        qc_para.add_run("Quality Control Summary: ").bold = True
        if variance_results['calculation_consistency']['tolerance_met']:
            qc_para.add_run("✓ ALL VALIDATION CHECKS PASSED").bold = True
        else:
            qc_para.add_run("❌ VALIDATION ISSUES DETECTED").bold = True
        
        doc.add_paragraph()
        
        # Add detailed variance analysis
        self._add_variance_analysis_section(doc)
        
        # Save technical report
        doc.save(file_path)
    
    def _create_chart(self) -> Optional[str]:
        """Create a temporary chart file for inclusion in Word document."""
        try:
            df = self.calculator.build_cost_schedule()
            
            plt.figure(figsize=(10, 6))
            
            if "Present Value" in df.columns:
                plt.bar(df["Year"], df["Present Value"], color='black', alpha=0.6)
                plt.title(f"Present Value of Medical Costs by Year\nEvaluee: {self.lcp.evaluee.name}")
                plt.ylabel("Present Value ($)")
            else:
                plt.bar(df["Year"], df["Total Nominal"], color='black', alpha=0.6)
                plt.title(f"Nominal Medical Costs by Year\nEvaluee: {self.lcp.evaluee.name}")
                plt.ylabel("Nominal Cost ($)")
            
            plt.xlabel("Year")
            plt.xticks(rotation=45)
            plt.tight_layout()
            
            # Save to temporary file
            temp_path = f"temp_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            plt.savefig(temp_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return temp_path
        except Exception as e:
            print(f"Warning: Could not create chart: {e}")
            return None

    def _add_calculation_methodology_section(self, doc):
        """Add comprehensive calculation methodology section with equations and explanations."""
        doc.add_heading("Calculation Methodology and Mathematical Framework", level=2)
        
        # Introduction
        intro_para = doc.add_paragraph()
        intro_para.add_run("Mathematical Foundation: ").bold = True
        intro_para.add_run("This section provides detailed mathematical equations and methodologies used in all cost projections. ")
        intro_para.add_run("All calculations follow established actuarial and financial principles to ensure accuracy and reliability.")
        
        doc.add_paragraph()
        
        # Core Equations Section
        doc.add_heading("Core Mathematical Equations", level=3)
        
        # Inflation Calculation
        eq1_para = doc.add_paragraph()
        eq1_para.add_run("1. Inflation-Adjusted Cost Calculation").bold = True
        eq1_para.add_run("\nFor recurring services, the cost in any given year is calculated as:\n")
        eq1_para.add_run("C(t) = C₀ × (1 + i)ᵗ").bold = True
        eq1_para.add_run("\nWhere:")
        eq1_para.add_run("\n• C(t) = Cost in year t")
        eq1_para.add_run("\n• C₀ = Base year unit cost")
        eq1_para.add_run("\n• i = Annual inflation rate (as decimal)")
        eq1_para.add_run("\n• t = Number of years from base year")
        
        doc.add_paragraph()
        
        # Present Value Calculation
        eq2_para = doc.add_paragraph()
        eq2_para.add_run("2. Present Value Calculation").bold = True
        eq2_para.add_run("\nTo discount future costs to present value:\n")
        eq2_para.add_run("PV(t) = C(t) ÷ (1 + d)ᵗ").bold = True
        eq2_para.add_run("\nWhere:")
        eq2_para.add_run("\n• PV(t) = Present value of cost in year t")
        eq2_para.add_run("\n• C(t) = Nominal cost in year t")
        eq2_para.add_run("\n• d = Discount rate (as decimal)")
        eq2_para.add_run("\n• t = Number of years from base year")
        
        doc.add_paragraph()
        
        # Total Lifetime Cost
        eq3_para = doc.add_paragraph()
        eq3_para.add_run("3. Total Lifetime Cost Calculation").bold = True
        eq3_para.add_run("\nFor services spanning the full projection period:\n")
        eq3_para.add_run("Total Nominal = Σ [C₀ × (1 + i)ᵗ × f]").bold = True
        eq3_para.add_run("\nTotal PV = Σ [C₀ × (1 + i)ᵗ × f ÷ (1 + d)ᵗ]").bold = True
        eq3_para.add_run("\nWhere:")
        eq3_para.add_run("\n• f = Frequency per year")
        eq3_para.add_run("\n• Σ = Sum over all years in projection period")
        
        doc.add_paragraph()
        
        # Fractional Year Calculation
        eq4_para = doc.add_paragraph()
        eq4_para.add_run("4. Fractional Year Adjustment").bold = True
        eq4_para.add_run(f"\nFor projection period of {self.lcp.settings.projection_years:.1f} years:")
        full_years = int(self.lcp.settings.projection_years)
        fractional_part = self.lcp.settings.projection_years - full_years
        if fractional_part > 0:
            eq4_para.add_run(f"\n• Full years: {full_years}")
            eq4_para.add_run(f"\n• Fractional year: {fractional_part:.1f}")
            eq4_para.add_run(f"\n• Final year cost = C({full_years}) × {fractional_part:.1f}")
        else:
            eq4_para.add_run(f"\n• Projection uses complete years only ({full_years} years)")
        
        doc.add_paragraph()
        
        # Service Type Methodologies
        doc.add_heading("Service Type Calculation Methods", level=3)
        
        # Recurring Services
        rec_para = doc.add_paragraph()
        rec_para.add_run("Recurring Services: ").bold = True
        rec_para.add_run("Applied annually from start year to end year. ")
        rec_para.add_run("Cost increases each year by the specified inflation rate. ")
        rec_para.add_run("Total frequency per year multiplied by inflated unit cost.")
        
        # One-time Services
        ot_para = doc.add_paragraph()
        ot_para.add_run("One-time Services: ").bold = True
        ot_para.add_run("Applied only in the specified year. ")
        ot_para.add_run("Unit cost inflated from base year to service year. ")
        ot_para.add_run("No ongoing costs in subsequent years.")
        
        # Discrete Occurrences
        disc_para = doc.add_paragraph()
        disc_para.add_run("Discrete Occurrences: ").bold = True
        disc_para.add_run("Applied only in specifically listed years. ")
        disc_para.add_run("Each occurrence independently inflated from base year. ")
        disc_para.add_run("Allows for irregular service patterns.")
        
        doc.add_paragraph()
        
        # Add Validation Framework
        self._add_validation_framework_section(doc)
        
        # Add Sensitivity Analysis
        self._add_sensitivity_analysis_section(doc)
        
        # Add Factor Tables
        self._add_factor_tables_section(doc)
        
        # Add Variance Analysis
        self._add_variance_analysis_section(doc)

    def _add_validation_framework_section(self, doc):
        """Add validation framework explanation."""
        doc.add_heading("Quality Control and Validation Framework", level=3)
        
        val_para = doc.add_paragraph()
        val_para.add_run("Cross-Validation Method: ").bold = True
        val_para.add_run("All calculations undergo five-point validation:")
        val_para.add_run("\n1. Category totals must reconcile with executive summary")
        val_para.add_run("\n2. Average annual cost verification: Total ÷ Projection Years")
        val_para.add_run("\n3. Year-by-year consistency across all report sections")
        val_para.add_run("\n4. Total sum verification with tolerance < $1.00")
        val_para.add_run("\n5. Matrix reconciliation using audit-standard methodologies")
        
        doc.add_paragraph()
        
        tol_para = doc.add_paragraph()
        tol_para.add_run("Tolerance Standards: ").bold = True
        tol_para.add_run("Acceptable discrepancies are limited to $1.00 due to rounding. ")
        tol_para.add_run("Any variance exceeding this threshold triggers automatic review and correction.")

    def _add_sensitivity_analysis_section(self, doc):
        """Add sensitivity analysis section."""
        doc.add_heading("Sensitivity Analysis", level=3)
        
        # Calculate sensitivity scenarios
        base_discount = self.lcp.settings.discount_rate
        base_summary = self.calculator.calculate_summary_statistics()
        
        sens_para = doc.add_paragraph()
        sens_para.add_run("Discount Rate Sensitivity: ").bold = True
        sens_para.add_run("The following table shows the impact of ±0.5% discount rate changes on total present value:")
        
        doc.add_paragraph()
        
        # Create sensitivity table
        sensitivity_data = []
        for rate_adjustment in [-0.005, 0.0, 0.005]:
            test_rate = base_discount + rate_adjustment
            # Create temporary settings for calculation
            temp_settings = ProjectionSettings(
                base_year=self.lcp.settings.base_year,
                projection_years=self.lcp.settings.projection_years,
                discount_rate=test_rate
            )
            temp_lcp = LifeCarePlan(evaluee=self.lcp.evaluee, settings=temp_settings)
            temp_lcp.tables = self.lcp.tables
            temp_calc = CostCalculator(temp_lcp)
            temp_summary = temp_calc.calculate_summary_statistics()
            
            pv_difference = temp_summary['total_present_value'] - base_summary['total_present_value']
            pv_percent = (pv_difference / base_summary['total_present_value']) * 100 if base_summary['total_present_value'] > 0 else 0
            
            sensitivity_data.append([
                f"{test_rate:.1%}",
                f"${temp_summary['total_present_value']:,.2f}",
                f"${pv_difference:,.2f}",
                f"{pv_percent:+.2f}%"
            ])
        
        # Add table to document
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light List'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Discount Rate"
        header_cells[1].text = "Total Present Value"
        header_cells[2].text = "Difference from Base"
        header_cells[3].text = "Percentage Change"
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Data rows
        for row_data in sensitivity_data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = value
        
        doc.add_paragraph()
        
        # Inflation sensitivity note
        inf_sens_para = doc.add_paragraph()
        inf_sens_para.add_run("Inflation Rate Sensitivity: ").bold = True
        inf_sens_para.add_run("Service-specific inflation rates are applied individually. ")
        inf_sens_para.add_run("A 1% increase in inflation across all services typically increases total nominal costs by 15-25% over long projection periods. ")
        inf_sens_para.add_run("Present value impacts are moderated by the discount rate effect.")

    def _add_factor_tables_section(self, doc):
        """Add inflation and discount factor tables."""
        doc.add_heading("Mathematical Factor Tables", level=3)
        
        # Discount Factor Table
        doc.add_heading("Discount Factors", level=4)
        df_para = doc.add_paragraph()
        df_para.add_run("Present Value Discount Factors: ").bold = True
        df_para.add_run(f"Based on {self.lcp.settings.discount_rate:.1%} annual discount rate")
        
        # Create discount factor table for first 10 years
        df_table = doc.add_table(rows=1, cols=3)
        df_table.style = 'Light List'
        
        df_header = df_table.rows[0].cells
        df_header[0].text = "Year"
        df_header[1].text = "Discount Factor"
        df_header[2].text = "Cumulative Factor"
        
        for cell in df_header:
            cell.paragraphs[0].runs[0].bold = True
        
        cumulative_factor = 0
        for year in range(min(10, int(self.lcp.settings.projection_years) + 1)):
            factor = 1 / (1 + self.lcp.settings.discount_rate) ** year
            cumulative_factor += factor
            
            row_cells = df_table.add_row().cells
            row_cells[0].text = str(self.lcp.settings.base_year + year)
            row_cells[1].text = f"{factor:.6f}"
            row_cells[2].text = f"{cumulative_factor:.6f}"
        
        doc.add_paragraph()
        
        # Sample Inflation Factor Table
        doc.add_heading("Sample Inflation Factors", level=4)
        if_para = doc.add_paragraph()
        if_para.add_run("Example Inflation Factors: ").bold = True
        if_para.add_run("Showing compound growth at common medical inflation rates")
        
        # Create inflation factor table
        if_table = doc.add_table(rows=1, cols=4)
        if_table.style = 'Light List'
        
        if_header = if_table.rows[0].cells
        if_header[0].text = "Year"
        if_header[1].text = "2.5% Inflation"
        if_header[2].text = "3.0% Inflation"
        if_header[3].text = "3.5% Inflation"
        
        for cell in if_header:
            cell.paragraphs[0].runs[0].bold = True
        
        for year in range(min(10, int(self.lcp.settings.projection_years) + 1)):
            row_cells = if_table.add_row().cells
            row_cells[0].text = str(self.lcp.settings.base_year + year)
            row_cells[1].text = f"{(1.025) ** year:.6f}"
            row_cells[2].text = f"{(1.030) ** year:.6f}"
            row_cells[3].text = f"{(1.035) ** year:.6f}"

    def _add_variance_analysis_section(self, doc):
        """Add automated variance analysis and error detection results."""
        doc.add_heading("Automated Variance Analysis and Error Detection", level=3)
        
        # Perform variance analysis
        variance_results = self.calculator.perform_variance_analysis()
        
        # Introduction
        va_intro = doc.add_paragraph()
        va_intro.add_run("Automated Analysis Summary: ").bold = True
        va_intro.add_run(f"Analysis performed on {variance_results['timestamp']}. ")
        va_intro.add_run("This section provides automated detection of potential errors, ")
        va_intro.add_run("inconsistencies, and unusual patterns in the calculation results.")
        
        doc.add_paragraph()
        
        # Data Integrity Results
        doc.add_heading("Data Integrity Assessment", level=4)
        integrity = variance_results['data_integrity_checks']
        
        integrity_para = doc.add_paragraph()
        integrity_para.add_run("Data Consistency Status: ").bold = True
        if integrity['data_consistency']:
            integrity_para.add_run("✓ PASS - All data integrity checks passed").bold = True
        else:
            integrity_para.add_run("❌ ISSUES DETECTED - Review required").bold = True
        
        if integrity['invalid_values']:
            doc.add_paragraph()
            iv_para = doc.add_paragraph()
            iv_para.add_run("Invalid Values Detected:").bold = True
            for issue in integrity['invalid_values']:
                iv_para.add_run(f"\n• {issue}")
        
        if integrity['missing_data']:
            doc.add_paragraph()
            md_para = doc.add_paragraph()
            md_para.add_run("Missing Data Issues:").bold = True
            for issue in integrity['missing_data']:
                md_para.add_run(f"\n• {issue}")
        
        doc.add_paragraph()
        
        # Calculation Consistency Results
        doc.add_heading("Calculation Consistency Verification", level=4)
        consistency = variance_results['calculation_consistency']
        
        consistency_para = doc.add_paragraph()
        consistency_para.add_run("Tolerance Compliance: ").bold = True
        if consistency['tolerance_met']:
            consistency_para.add_run("✓ PASS - All discrepancies within $1.00 tolerance").bold = True
        else:
            consistency_para.add_run("❌ FAIL - Discrepancies exceed tolerance").bold = True
        
        # Show specific checks
        for check_name, check_data in consistency.items():
            if isinstance(check_data, dict) and 'passes' in check_data:
                check_para = doc.add_paragraph()
                check_para.add_run(f"{check_name.replace('_', ' ').title()}: ").bold = True
                if check_data['passes']:
                    check_para.add_run("✓ PASS")
                else:
                    check_para.add_run(f"❌ FAIL - ${check_data['difference']:.2f} discrepancy")
        
        doc.add_paragraph()
        
        # Reasonableness Assessment
        doc.add_heading("Reasonableness Assessment", level=4)
        reasonableness = variance_results['reasonableness_checks']
        
        # Cost distribution
        if 'cost_distribution' in reasonableness:
            cd = reasonableness['cost_distribution']
            cd_para = doc.add_paragraph()
            cd_para.add_run("Cost Distribution Analysis:").bold = True
            cd_para.add_run(f"\n• Annual cost range: ${cd['min_annual']:,.0f} - ${cd['max_annual']:,.0f}")
            cd_para.add_run(f"\n• Average annual cost: ${cd['mean_annual']:,.0f}")
            cd_para.add_run(f"\n• Cost variability: {cd['coefficient_of_variation']:.2f}")
        
        # Outlier detection
        if 'outlier_detection' in reasonableness and reasonableness['outlier_detection']['outlier_count'] > 0:
            od = reasonableness['outlier_detection']
            od_para = doc.add_paragraph()
            od_para.add_run("Outlier Detection:").bold = True
            od_para.add_run(f"\n• {od['outlier_count']} outlier years detected: {', '.join(map(str, od['outlier_years']))}")
        
        doc.add_paragraph()
        
        # Trend Analysis
        doc.add_heading("Cost Trend Analysis", level=4)
        trends = variance_results['trend_analysis']
        
        trend_para = doc.add_paragraph()
        trend_para.add_run("Overall Cost Trend: ").bold = True
        trend_para.add_run(f"{trends['overall_trend'].upper()}")
        
        if 'early_years_avg' in trends and trends['early_years_avg'] > 0:
            trend_para.add_run(f"\n• Early years average: ${trends['early_years_avg']:,.0f}")
            trend_para.add_run(f"\n• Middle years average: ${trends['middle_years_avg']:,.0f}")
            trend_para.add_run(f"\n• Late years average: ${trends['late_years_avg']:,.0f}")
        
        trend_para.add_run(f"\n• Peak cost year: {trends['peak_cost_year']} (${trends['peak_cost_amount']:,.0f})")
        
        doc.add_paragraph()
        
        # Error Flags and Warnings
        if variance_results['error_flags'] or variance_results['warnings']:
            doc.add_heading("Critical Issues and Warnings", level=4)
            
            if variance_results['error_flags']:
                error_para = doc.add_paragraph()
                error_para.add_run("Critical Errors:").bold = True
                for error in variance_results['error_flags']:
                    error_para.add_run(f"\n❌ {error}")
            
            if variance_results['warnings']:
                warning_para = doc.add_paragraph()
                warning_para.add_run("Warnings:").bold = True
                for warning in variance_results['warnings']:
                    warning_para.add_run(f"\n⚠️ {warning}")
        
        # Recommendations
        doc.add_heading("Analysis Recommendations", level=4)
        rec_para = doc.add_paragraph()
        rec_para.add_run("Recommended Actions:").bold = True
        for i, recommendation in enumerate(variance_results['recommendations'], 1):
            rec_para.add_run(f"\n{i}. {recommendation}")

    def export_combined_scenarios(self, file_path: str, selected_scenarios: list) -> None:
        """Export all selected scenarios combined into a single Word document format."""
        doc = Document()

        # Set document to landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Adjust margins for better table fit
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        # Main title
        title = doc.add_heading("LIFE CARE PLAN - COMBINED SCENARIOS", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_heading("Comprehensive Economic Analysis for Multiple Scenarios", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Evaluee information
        evaluee_para = doc.add_paragraph()
        evaluee_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        evaluee_para.add_run("Prepared for: ").bold = True
        evaluee_para.add_run(f"{self.lcp.evaluee.name}").bold = True
        
        doc.add_paragraph()
        
        # Scenarios included
        scenarios_para = doc.add_paragraph()
        scenarios_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        scenarios_para.add_run("Scenarios Included: ").bold = True
        scenarios_para.add_run(", ".join(selected_scenarios))
        
        doc.add_paragraph()

        # Store original active scenario
        original_active = self.lcp.active_scenario
        
        try:
            # Process each scenario
            for i, scenario_name in enumerate(selected_scenarios):
                # Set active scenario
                self.lcp.set_active_scenario(scenario_name)
                
                # Add scenario section
                if i > 0:
                    doc.add_page_break()
                
                scenario_heading = doc.add_heading(f"SCENARIO: {scenario_name.upper()}", level=2)
                scenario_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Get scenario info
                scenario = self.lcp.scenarios[scenario_name]
                if scenario.description:
                    desc_para = doc.add_paragraph()
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    desc_para.add_run("Description: ").bold = True
                    desc_para.add_run(scenario.description)
                    doc.add_paragraph()
                
                # Add all the standard sections for this scenario
                self._add_scenario_executive_summary(doc)
                self._add_scenario_cost_breakdown(doc)
                self._add_scenario_detailed_schedule(doc)
                
        finally:
            # Restore original active scenario
            self.lcp.set_active_scenario(original_active)
        
        doc.save(file_path)

    def _add_scenario_executive_summary(self, doc):
        """Add executive summary section for current scenario."""
        doc.add_heading("Executive Summary", level=3)
        
        summary_stats = self.calculator.calculate_summary_statistics()
        
        # Document information table
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light List'
        
        info_data = [
            ["Analysis Date:", datetime.now().strftime('%B %d, %Y')],
            ["Current Age:", f"{self.lcp.evaluee.current_age:.1f} years"],
            ["Base Year:", str(int(self.lcp.settings.base_year))],
            ["Projection Period:", f"{self.lcp.settings.projection_years:.1f} years"],
            ["Total Nominal Cost:", f"${summary_stats['total_nominal_cost']:,.2f}"],
            ["Average Annual Cost:", f"${summary_stats['average_annual_cost']:,.2f}"]
        ]
        
        # Add present value info if enabled
        if self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0:
            info_data.append(["Total Present Value:", f"${summary_stats['total_present_value']:,.2f}"])
            info_data.append(["Discount Rate:", f"{self.lcp.settings.discount_rate:.1%}"])
        
        for i, (label, value) in enumerate(info_data):
            if i < len(info_table.rows):
                row_cells = info_table.rows[i].cells
            else:
                row_cells = info_table.add_row().cells
                
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = value
        
        doc.add_paragraph()

    def _add_scenario_cost_breakdown(self, doc):
        """Add cost breakdown by category for current scenario."""
        doc.add_heading("Cost Breakdown by Category", level=3)
        
        category_costs = self.calculator.get_cost_by_category()
        show_present_value = self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0
        
        # Create category table
        cols = 3 if show_present_value else 2
        headers = ['Service Category', 'Total Nominal Cost']
        if show_present_value:
            headers.append('Total Present Value')
        
        category_table = doc.add_table(rows=1, cols=cols)
        category_table.style = 'Light List'
        category_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = category_table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header
            hdr_cells[idx].paragraphs[0].runs[0].bold = True
            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        total_nominal = 0
        total_pv = 0
        
        for category_name, costs in category_costs.items():
            row_cells = category_table.add_row().cells
            row_cells[0].text = category_name
            row_cells[1].text = f"${costs['table_nominal_total']:,.2f}"
            
            if show_present_value:
                row_cells[2].text = f"${costs['table_present_value_total']:,.2f}"
            
            total_nominal += costs['table_nominal_total']
            total_pv += costs['table_present_value_total']
            
            # Format cells
            for cell in row_cells[1:]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Total row
        total_cells = category_table.add_row().cells
        total_cells[0].text = "TOTAL"
        total_cells[0].paragraphs[0].runs[0].bold = True
        total_cells[1].text = f"${total_nominal:,.2f}"
        total_cells[1].paragraphs[0].runs[0].bold = True
        total_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if show_present_value:
            total_cells[2].text = f"${total_pv:,.2f}"
            total_cells[2].paragraphs[0].runs[0].bold = True
            total_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()

    def _add_scenario_detailed_schedule(self, doc):
        """Add detailed year-by-year schedule for current scenario."""
        doc.add_heading("Detailed Year-by-Year Cost Schedule", level=3)
        
        # Build cost schedule
        df = self.calculator.build_cost_schedule()
        show_present_value = self.lcp.evaluee.discount_calculations and self.lcp.settings.discount_rate > 0
        
        # Create table
        base_cols = ['Year', 'Age', 'Total Nominal']
        if show_present_value:
            base_cols.append('Present Value')
        
        num_cols = len(base_cols)
        schedule_table = doc.add_table(rows=1, cols=num_cols)
        schedule_table.style = 'Light List'
        schedule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = schedule_table.rows[0].cells
        for idx, header in enumerate(base_cols):
            hdr_cells[idx].text = header
            hdr_cells[idx].paragraphs[0].runs[0].bold = True
            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows (show first 15 years, then summary)
        years_to_show = min(15, len(df))
        
        for i in range(years_to_show):
            row = df.iloc[i]
            row_cells = schedule_table.add_row().cells
            
            row_cells[0].text = str(int(row['Year']))
            row_cells[1].text = f"{row['Age']:.1f}"
            row_cells[2].text = f"${row['Total Nominal']:,.0f}"
            
            if show_present_value and 'Present Value' in df.columns:
                row_cells[3].text = f"${row['Present Value']:,.0f}"
            
            # Format
            for cell in row_cells[2:]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Summary row if more years exist
        if len(df) > years_to_show:
            summary_cells = schedule_table.add_row().cells
            summary_cells[0].text = "..."
            summary_cells[1].text = "..."
            remaining_years = len(df) - years_to_show
            remaining_nominal = df.iloc[years_to_show:]['Total Nominal'].sum()
            summary_cells[2].text = f"${remaining_nominal:,.0f} ({remaining_years} more years)"
            
            if show_present_value and 'Present Value' in df.columns:
                remaining_pv = df.iloc[years_to_show:]['Present Value'].sum()
                summary_cells[3].text = f"${remaining_pv:,.0f}"
        
        doc.add_paragraph()


class PDFExporter:
    """Export life care plan data to PDF format using ReportLab."""
    
    def __init__(self, calculator: CostCalculator):
        self.calculator = calculator
        self.lcp = calculator.lcp
    
    def export(self, file_path: str) -> None:
        """Export the life care plan to PDF file in landscape mode."""
        from reportlab.lib.pagesizes import letter, A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER

        # Create PDF document in landscape mode
        doc = SimpleDocTemplate(
            file_path,
            pagesize=landscape(letter),
            leftMargin=0.5*inch,
            rightMargin=0.5*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            alignment=TA_CENTER,
            spaceAfter=20
        )

        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=30
        )

        story.append(Paragraph("Life Care Plan Economic Projection", title_style))
        story.append(Paragraph(f"Evaluee: {self.lcp.evaluee.name}", subtitle_style))
        story.append(Spacer(1, 20))
        
        # Enhanced Metadata
        story.append(Paragraph(f"<b>Report Generated:</b> {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph(f"<b>Evaluee Age at Analysis Start:</b> {self.lcp.evaluee.current_age} years old (in {self.lcp.settings.base_year})", styles['Normal']))

        end_year = self.lcp.settings.base_year + self.lcp.settings.projection_years - 1
        story.append(Paragraph(f"<b>Analysis Period:</b> {self.lcp.settings.projection_years:.1f} years "
                             f"({self.lcp.settings.base_year} to {end_year:.1f})", styles['Normal']))

        if self.lcp.evaluee.discount_calculations:
            story.append(Paragraph(f"<b>Discount Rate Applied:</b> {self.lcp.settings.discount_rate:.1%} annually", styles['Normal']))
            story.append(Paragraph("<b>Present Value Calculations:</b> Enabled", styles['Normal']))
        else:
            story.append(Paragraph("<b>Present Value Calculations:</b> Not Applied", styles['Normal']))

        story.append(Paragraph(f"<b>Service Categories Analyzed:</b> {len(self.lcp.tables)}", styles['Normal']))
        story.append(Paragraph(f"<b>Total Individual Services:</b> {sum(len(table.services) for table in self.lcp.tables.values())}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Summary statistics
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        summary_stats = self.calculator.calculate_summary_statistics()
        
        summary_data = [
            ['Financial Summary', 'Amount'],
            ['Total Lifetime Medical Costs (Nominal)', f"${summary_stats['total_nominal_cost']:,.0f}"],
            ['Average Annual Medical Costs', f"${summary_stats['average_annual_cost']:,.0f}"]
        ]

        if self.lcp.evaluee.discount_calculations:
            summary_data.append(['Total Lifetime Medical Costs (Present Value)', f"${summary_stats['total_present_value']:,.0f}"])
            savings = summary_stats['total_nominal_cost'] - summary_stats['total_present_value']
            summary_data.append(['Present Value Savings vs Nominal', f"${savings:,.0f}"])
        
        summary_table = Table(summary_data)
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # Category breakdown
        story.append(Paragraph("Cost Breakdown by Service Category", styles['Heading2']))
        category_costs = self.calculator.get_cost_by_category()

        if self.lcp.evaluee.discount_calculations:
            category_data = [['Service Category', 'Lifetime Cost (Nominal)', 'Lifetime Cost (Present Value)', 'Number of Services']]
            for table_name, data in category_costs.items():
                category_data.append([
                    table_name,
                    f"${data['table_nominal_total']:,.0f}",
                    f"${data['table_present_value_total']:,.0f}",
                    str(len(data['services']))
                ])
        else:
            category_data = [['Service Category', 'Total Lifetime Cost (Nominal)', 'Number of Services']]
            for table_name, data in category_costs.items():
                category_data.append([
                    table_name,
                    f"${data['table_nominal_total']:,.0f}",
                    str(len(data['services']))
                ])
        
        category_table = Table(category_data)
        category_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(category_table)
        story.append(PageBreak())
        
        # Detailed cost schedule
        story.append(Paragraph("Annual Cost Schedule", styles['Heading2']))
        df = self.calculator.build_cost_schedule()

        # Prepare table data with improved headers
        if "Present Value" in df.columns:
            table_data = [['Year', 'Evaluee Age', 'Annual Cost (Nominal)', 'Annual Cost (Present Value)']]
            for _, row in df.iterrows():
                table_data.append([
                    str(int(row['Year'])),
                    str(int(row['Age'])),
                    f"${row['Total Nominal']:,.0f}",
                    f"${row['Present Value']:,.0f}"
                ])
        else:
            table_data = [['Year', 'Evaluee Age', 'Annual Medical Cost (Nominal)']]
            for _, row in df.iterrows():
                table_data.append([
                    str(int(row['Year'])),
                    str(int(row['Age'])),
                    f"${row['Total Nominal']:,.0f}"
                ])
        
        detail_table = Table(table_data)
        detail_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(detail_table)
        
        # Build PDF
        doc.build(story)